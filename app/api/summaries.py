from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session, joinedload
from pydantic import BaseModel

from app.database import get_db
from app.models import WeeklySummary, WeekPeriod
from app.schemas import SummaryOut, SummaryUpdate
from app.services.report_service import get_or_create_current_period
from app.services.summary_service import generate_summary, get_summary_prompt, save_summary_prompt, get_system_prompt, save_system_prompt

from app.api.deps import get_current_user

router = APIRouter(prefix="/api/summaries", tags=["汇总管理"], dependencies=[Depends(get_current_user)])

class PromptUpdate(BaseModel):
    user_template: str
    system_prompt: str

@router.get("/prompt")
def get_prompt():
    """获取当前的周报汇总模版配置（包括用户模板和系统角色提示词）。

    Returns:
        包含 user_template 和 system_prompt 的字典。
    """
    return {
        "user_template": get_summary_prompt(),
        "system_prompt": get_system_prompt()
    }

@router.put("/prompt")
def update_prompt(data: PromptUpdate):
    """更新并保存周报汇总模板配置。

    Args:
        data: 包含更新内容的 PromptUpdate 数据传输对象。

    Returns:
        包含成功提示的字典。
    """
    save_summary_prompt(data.user_template)
    save_system_prompt(data.system_prompt)
    return {"message": "提示词配置已更新"}


def load_docx_styles():
    """从本地 JSON 配置文件中加载 DOCX 排版样式。

    若配置文件不存在或解析异常，则使用默认的公文排版参数进行兜底。

    Returns:
        包含各级标题和正文样式属性的字典。
    """
    import os
    import json

    # 默认兜底配置
    default_styles = {
        "heading_1": {
            "ch_font": "方正小标宋简体",
            "ch_size": 22,
            "en_font": "Times New Roman",
            "en_size": 10
        },
        "heading_2": {
            "ch_font": "黑体",
            "ch_size": 16,
            "en_font": "Times New Roman",
            "en_size": 10
        },
        "heading_3": {
            "ch_font": "楷体_GB2312",
            "ch_size": 16,
            "en_font": "Times New Roman",
            "en_size": 10
        },
        "body": {
            "ch_font": "仿宋_GB2312",
            "ch_size": 16,
            "en_font": "Times New Roman",
            "en_size": 10
        }
    }

    config_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docx_style_config.json")
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                styles = json.load(f)
                # 针对每一项补齐缺省的 key，防止部分节点配置缺失导致报错
                for key, default_val in default_styles.items():
                    if key not in styles:
                        styles[key] = default_val
                    else:
                        for sub_key, sub_val in default_val.items():
                            if sub_key not in styles[key]:
                                styles[key][sub_key] = sub_val
                return styles
        except Exception:
            pass

    return default_styles


def build_docx_document(summary: WeeklySummary):
    """根据 WeeklySummary 的 Markdown 文本内容以及本地样式配置构建 Document。

    Args:
        summary: WeeklySummary 实例。

    Returns:
        Document: 生成的 Word 文档对象。
    """
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    styles = load_docx_styles()

    def apply_fonts(run, style_cfg):
        """设置 Word 运行段的中西文字体及字号。"""
        run.font.color.rgb = RGBColor(0, 0, 0)
        ch_font = style_cfg["ch_font"]
        ch_size = style_cfg["ch_size"]
        en_font = style_cfg.get("en_font", "Times New Roman")
        
        # 1. 设置中西文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)
        rFonts.set(qn('w:eastAsia'), ch_font)
        
        # 2. 统一设置字号
        run.font.size = Pt(ch_size)

    for line in summary.summary_content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                apply_fonts(r, styles["heading_1"])
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            for r in p.runs:
                apply_fonts(r, styles["heading_2"])
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            for r in p.runs:
                apply_fonts(r, styles["heading_3"])
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parts = line[2:].split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
                apply_fonts(run, styles["body"])
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
                apply_fonts(run, styles["body"])
                
    return doc


def auto_send_period_summary(db: Session, period: WeekPeriod, email: str) -> bool:
    """对目标统计周期执行最新性比对，自动生成 AI 汇总报告（若有新提交），
    并将其导出为 docx 附件发送到指定邮箱。发送成功后清理历史旧版本汇总。

    Args:
        db: 数据库会话对象。
        period: 统计周期实体。
        email: 接收邮箱地址。

    Returns:
        bool: 是否发送成功。
    """
    from loguru import logger
    import tempfile
    import os
    import sys
    from app.models import WeeklyReport, WeeklySummary
    from app.services.summary_service import generate_summary

    # 1. 检测最新性并自动重构生成
    latest_report = (
        db.query(WeeklyReport)
        .filter(WeeklyReport.week_period_id == period.id)
        .order_by(WeeklyReport.submitted_at.desc())
        .first()
    )
    if not latest_report:
        logger.warning(f"周期 {period.id} ({period.week_start} 至 {period.week_end}) 暂无任何成员提交周报，跳过自动汇总发送。")
        return False

    latest_summary = (
        db.query(WeeklySummary)
        .filter(WeeklySummary.week_period_id == period.id)
        .order_by(WeeklySummary.generated_at.desc())
        .first()
    )

    summary_to_send = None
    
    if not latest_summary or latest_report.submitted_at > latest_summary.generated_at:
        logger.info(f"周期 {period.id} 检测到新提交或尚未生成过汇总，开始自动生成最新汇总报告...")
        try:
            summary_to_send = generate_summary(db, period)
        except Exception as e:
            logger.error(f"周期 {period.id} 自动生成最新汇总失败: {e}", exc_info=True)
            return False
    else:
        logger.info(f"周期 {period.id} 无新成员提交，将直接使用历史最新版本进行发送。")
        summary_to_send = latest_summary

    # 2. 生成 DOCX 文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = tmp.name
        doc = build_docx_document(summary_to_send)
        doc.save(temp_path)
    
    start_str = period.week_start.strftime("%Y%m%d")
    end_str = period.week_end.strftime("%Y%m%d")

    # 3. 投递邮件
    try:
        try:
            from send_mail import send_mail
        except ImportError:
            root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
            if root_dir not in sys.path:
                sys.path.append(root_dir)
            from send_mail import send_mail

        subject = f"自动发送：工作周报-数据库团队_{start_str}-{end_str}"
        body = f"您好，当前周期 ({period.week_start} 至 {period.week_end}) 的数据库团队工作周报已自动截止并生成最新汇总，请查收附件。"
        
        send_mail(
            subject=subject,
            body=body,
            to=email,
            attach=[temp_path]
        )
        logger.success(f"周期 {period.id} 自动发送汇总邮件成功！收件人: {email}")
    except Exception as e:
        logger.error(f"周期 {period.id} 自动发送邮件投递失败: {e}", exc_info=True)
        return False
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass

    # 4. 邮件发送成功后，清理当前周期的历史旧汇总记录，仅保留发送成功的目标版本
    try:
        db.query(WeeklySummary).filter(
            WeeklySummary.week_period_id == period.id,
            WeeklySummary.id != summary_to_send.id
        ).delete()
        db.commit()
        logger.info(f"周期 {period.id} 历史旧草稿清理完毕，仅保留最终发送的 Summary ID: {summary_to_send.id}")
    except Exception as e:
        logger.error(f"周期 {period.id} 清理历史旧汇总记录时异常: {e}", exc_info=True)

    return True


async def auto_send_summary_loop():
    """后台定时检查截止时间及延时，并自动发送的轮询协程守护任务。"""
    import asyncio
    from datetime import datetime, timedelta
    from loguru import logger
    from app.database import SessionLocal
    from app.config import get_deadline_config
    from app.services.report_service import get_or_create_current_period

    logger.info("自动发送周报汇总后台轮询协程启动就绪。")
    while True:
        try:
            await asyncio.sleep(60)

            cfg = get_deadline_config()
            if not cfg.auto_send_enabled or not cfg.auto_send_email:
                continue

            db = SessionLocal()
            try:
                period = get_or_create_current_period(db)
                if not period:
                    continue

                if period.auto_sent_at is not None:
                    continue

                send_threshold = period.deadline + timedelta(minutes=period.auto_send_delay)
                if datetime.now() < send_threshold:
                    continue

                logger.info(
                    f"截止时间延时已到达阈值！开始对周期 {period.id} 执行自动汇总与投递流程。收件人: {cfg.auto_send_email}"
                )
                success = auto_send_period_summary(db, period, cfg.auto_send_email)
                if success:
                    period.auto_sent_at = datetime.now()
                    db.commit()
                    logger.success(f"周期 {period.id} 状态已成功更新 auto_sent_at 时间标记。")
            except Exception as err:
                logger.error(f"自动发送轮询内部事务出现故障: {err}", exc_info=True)
            finally:
                db.close()

        except asyncio.CancelledError:
            logger.info("自动发送周报汇总协程收到取消请求，退出轮询。")
            break
        except Exception as e:
            logger.error(f"自动发送后台协程发生主异常: {e}", exc_info=True)


@router.post("/generate", response_model=SummaryOut)
def trigger_summary(db: Session = Depends(get_db)):
    """手动触发生成当前周期的 AI 工作周报汇总。

    若当前周期没有任何成员提交周报，则拒绝生成。

    Args:
        db: 数据库 Session 对象。

    Returns:
        新生成的 WeeklySummary 汇总实体。

    Raises:
        HTTPException: 当无周报可汇总导致生成失败时抛出 400。
    """
    period = get_or_create_current_period(db)
    try:
        summary = generate_summary(db, period)
    except ValueError as e:
        raise HTTPException(400, str(e))
    return summary


@router.get("", response_model=list[SummaryOut])
def list_summaries(db: Session = Depends(get_db)):
    """获取所有历史周报 AI 汇总结果列表。

    默认以生成时间倒序排列。

    Args:
        db: 数据库 Session 对象。

    Returns:
        WeeklySummary 汇总对象列表。
    """
    return (
        db.query(WeeklySummary)
        .options(joinedload(WeeklySummary.week_period))
        .order_by(WeeklySummary.generated_at.desc())
        .all()
    )


@router.get("/{summary_id}", response_model=SummaryOut)
def get_summary(summary_id: int, db: Session = Depends(get_db)):
    """获取指定 ID 的 AI 汇总记录详情。

    Args:
        summary_id: 汇总记录的 ID。
        db: 数据库 Session 对象。

    Returns:
        WeeklySummary 实体详情。

    Raises:
        HTTPException: 当寻找的汇总不存在时报 404。
    """
    summary = (
        db.query(WeeklySummary)
        .options(joinedload(WeeklySummary.week_period))
        .filter(WeeklySummary.id == summary_id)
        .first()
    )
    if not summary:
        raise HTTPException(404, "汇总不存在")
    return summary


@router.put("/{summary_id}", response_model=SummaryOut)
def update_summary(summary_id: int, data: SummaryUpdate, db: Session = Depends(get_db)):
    """编辑更新指定 ID 的 AI 汇总报告文本内容。

    Args:
        summary_id: 汇总记录的 ID。
        data: 修改内容的数据对象。
        db: 数据库 Session 对象。

    Returns:
        更新后的 WeeklySummary 实体对象。

    Raises:
        HTTPException: 当汇总不存在时抛出 404。
    """
    summary = db.get(WeeklySummary, summary_id)
    if not summary:
        raise HTTPException(404, "汇总不存在")
    summary.summary_content = data.summary_content
    db.commit()
    db.refresh(summary)
    return summary


@router.delete("/{summary_id}")
def delete_summary(summary_id: int, db: Session = Depends(get_db)):
    """删除指定的 AI 汇总记录。

    只允许删除当前周期的汇总，历史周期的汇总被保护且不允许被删除。

    Args:
        summary_id: 汇总记录的 ID。
        db: 数据库 Session 对象。

    Returns:
        包含成功提示的字典。

    Raises:
        HTTPException: 当汇总不存在或试图删除历史周期汇总时抛出。
    """
    summary = db.get(WeeklySummary, summary_id)
    if not summary:
        raise HTTPException(404, "汇总不存在")
    current_period = get_or_create_current_period(db)
    if summary.week_period_id != current_period.id:
        raise HTTPException(403, "历史周期的汇总不允许删除")
    db.delete(summary)
    db.commit()
    return {"message": "已删除"}


@router.get("/{summary_id}/download")
def download_summary(summary_id: int, db: Session = Depends(get_db)):
    """将指定 ID 的 AI 汇总报告导出并下载为 Word 格式（.docx）文档。

    从数据库中提取周报汇总 Markdown 文本，逐行解析其标题和列表结构，
    并使用 python-docx 动态排版、应用中西文字体（Times New Roman + 仿宋），
    最终以二进制文件流的形式返回给客户端。

    Args:
        summary_id: 汇总记录的 ID。
        db: 数据库 Session 对象。

    Returns:
        包含 .docx 文件流的 FastAPI Response 响应对象。

    Raises:
        HTTPException: 当汇总不存在时抛出 404。
    """
    import io
    import urllib.parse
    from fastapi.responses import Response

    summary = (
        db.query(WeeklySummary)
        .options(joinedload(WeeklySummary.week_period))
        .filter(WeeklySummary.id == summary_id)
        .first()
    )
    if not summary:
        raise HTTPException(404, "汇总不存在")
        
    doc = build_docx_document(summary)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    start_str = summary.week_period.week_start.strftime('%Y%m%d')
    end_str = summary.week_period.week_end.strftime('%Y%m%d')
    filename = f"工作周报-数据库团队_{start_str}-{end_str}.docx"
    encoded_filename = urllib.parse.quote(filename)
    
    return Response(
        content=file_stream.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"
        }
    )


class MailSendRequest(BaseModel):
    """邮件发送请求体模型。"""
    to: list[str]
    subject: str | None = None
    body: str | None = None


@router.post("/{summary_id}/send")
def send_summary_mail(
    summary_id: int,
    req: MailSendRequest,
    db: Session = Depends(get_db)
):
    """向指定收件人列表发送周报汇总报告邮件，并附带生成的 Word 格式文档。

    Args:
        summary_id: 需要发送的 AI 周报汇总记录 ID。
        req: 包含收件人、主题和正文的 MailSendRequest 对象。
        db: 数据库 Session 对象。

    Returns:
        包含成功提示信息和邮件网关返回的字典。

    Raises:
        HTTPException: 当汇总记录不存在或生成附件、发送邮件失败时抛出。
    """
    import os
    import sys
    import tempfile
    try:
        from send_mail import send_mail
    except ImportError:
        root_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        if root_dir not in sys.path:
            sys.path.append(root_dir)
        from send_mail import send_mail

    summary = (
        db.query(WeeklySummary)
        .options(joinedload(WeeklySummary.week_period))
        .filter(WeeklySummary.id == summary_id)
        .first()
    )
    if not summary:
        raise HTTPException(status_code=404, detail="汇总不存在")

    start_str = summary.week_period.week_start.strftime('%Y%m%d')
    end_str = summary.week_period.week_end.strftime('%Y%m%d')

    # 生成临时 Word 文件
    doc = build_docx_document(summary)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        temp_path = tmp.name

    subject = req.subject or f"工作周报 ({start_str} - {end_str})"
    body = req.body or f"附件为 {start_str}-{end_str} 的工作周报汇总，请查收。"

    # 调用邮件推送工具
    try:
        res = send_mail(
            subject=subject,
            body=body,
            to=req.to,
            attach=[temp_path]
        )
        return {
            "message": "邮件发送成功",
            "status_code": res.status_code,
            "detail": res.text
        }
    except Exception as e:
        from loguru import logger
        logger.error(f"邮件发送服务异常: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"邮件发送服务异常: {e}")
    finally:
        # 清理临时文件
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
